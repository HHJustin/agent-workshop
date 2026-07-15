"""
统一文档加载器 — 自动识别文件类型并路由到对应处理器

支持的格式：TXT / MD / PDF / DOCX
待扩展：PNG / JPG（图片 Caption 生成）

Author: 程响
"""

import re
from pathlib import Path
from typing import Optional

from langchain_core.documents import Document

from .pdf_handler import PDFHandler, get_pdf_complexity, analyze_pdf_structure

# 支持的文件扩展名
ALLOWED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".doc"}


def load_document(
    file_path: str | Path,
    pdf_fallback: str = "auto",
) -> list[Document]:
    """
    统一文档加载入口 — 自动识别文件类型

    Args:
        file_path: 文件路径
        pdf_fallback: PDF 处理策略

    Returns:
        list[Document]: LangChain Document 列表
    """
    file_path = Path(file_path).resolve()
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = file_path.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: '{ext}'。支持: {', '.join(sorted(ALLOWED_EXTENSIONS))}")

    if ext in (".txt", ".md", ".markdown"):
        return _load_text_robust(file_path)
    if ext in (".docx", ".doc"):
        return _load_docx(file_path)
    if ext == ".pdf":
        return PDFHandler(fallback=pdf_fallback).process(str(file_path))

    raise ValueError(f"无法处理: {file_path}")


def load_document_with_info(
    file_path: str | Path,
    pdf_fallback: str = "auto",
) -> tuple[list[Document], dict]:
    """加载文档 + 返回文件元信息"""
    documents = load_document(file_path, pdf_fallback)
    fp = Path(file_path)
    info = {
        "file_name": fp.name,
        "file_size_kb": round(fp.stat().st_size / 1024, 1),
        "file_type": fp.suffix.lower(),
        "document_count": len(documents),
        "total_chars": sum(len(d.page_content) for d in documents),
    }
    if fp.suffix.lower() == ".pdf":
        try:
            c = get_pdf_complexity(str(fp))
            info["pdf_complexity"] = {
                "page_count": c.page_count, "has_table": c.has_table,
                "has_multi_column": c.has_multi_column, "is_scanned": c.is_scanned,
                "used_mineru": c.needs_mineru,
            }
        except Exception:
            pass
    return documents, info


# ============================================================
# TXT：多编码容错 + 清洗
# ============================================================

def _read_txt_robust(file_path: Path) -> str:
    """读取 TXT 文件，依次尝试 UTF-8 → GBK → GB2312"""
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法解码文件: {file_path}")


def _clean_text(text: str) -> str:
    """清洗文本：去控制字符、合并空行、剔除页眉页脚标记"""
    # 删除不可见控制字符（保留换行、制表符）
    text = re.sub(r"[^\S\r\n\t\x20-\x7E一-鿿　-〿＀-￯]", "", text)
    # 合并多个空行
    text = re.sub(r"\n\s*\n", "\n\n", text)
    # 删除"第 X 页"类标记
    text = re.sub(r"第\s*\d+\s*页", "", text)
    return text.strip()


def _load_text_robust(file_path: Path) -> list[Document]:
    """加载 TXT/MD 文件（多编码容错 + 清洗）"""
    content = _read_txt_robust(file_path)
    content = _clean_text(content)

    if not content.strip():
        return []

    return [Document(
        page_content=content,
        metadata={
            "_source": str(file_path),
            "_loader": "TextLoader",
            "_file_name": file_path.name,
            "_extension": file_path.suffix.lower(),
        },
    )]


# ============================================================
# Word (.docx)：结构化提取
# ============================================================

def _load_docx(file_path: Path) -> list[Document]:
    """
    加载 Word 文档，保留标题层级和表格结构

    处理策略：
    - 标题 → 映射为 Markdown #/##/###
    - 正文 → 保留原文
    - 表格 → 转 Markdown 表格，独立成 chunk
    - 图片 → 提取 caption（题注），OCR 后续迭代
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        markdown_lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 标题映射为 Markdown
            if para.style and para.style.name and "Heading" in para.style.name:
                try:
                    level = int(para.style.name.replace("Heading", "").strip())
                    level = min(level, 3)  # 最多到 ###
                    markdown_lines.append(f"\n{'#' * level} {text}\n")
                except ValueError:
                    markdown_lines.append(text)
            else:
                markdown_lines.append(text)

        # 提取表格
        tables_md = []
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if r_idx == 0:
                    rows.append("|" + "|".join(["---"] * len(cells)) + "|")
            if rows:
                tables_md.append(f"\n【表格 {t_idx + 1}】\n" + "\n".join(rows))

        # 正文 + 表格合并
        full_text = "\n".join(markdown_lines)
        if tables_md:
            full_text += "\n\n" + "\n\n".join(tables_md)

        if not full_text.strip():
            return []

        return [Document(
            page_content=full_text,
            metadata={
                "_source": str(file_path),
                "_loader": "DocxLoader",
                "_file_name": file_path.name,
                "_extension": ".docx",
                "_tables": len(doc.tables),
            },
        )]

    except ImportError:
        raise ImportError("python-docx 需要安装: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Word 文档解析失败: {e}") from e


# ============================================================
# 自测
# ============================================================

if __name__ == "__main__":
    import sys
    import logging
    logging.basicConfig(level=logging.INFO)

    if len(sys.argv) < 2:
        print("用法: python loader.py <file_path>")
        sys.exit(1)

    docs, info = load_document_with_info(sys.argv[1])
    print(f"\n文件信息: {info}")
    print(f"\n文档数量: {len(docs)}")
    for i, doc in enumerate(docs[:3]):
        print(f"\n--- Document {i+1} ---")
        print(doc.page_content[:300])
